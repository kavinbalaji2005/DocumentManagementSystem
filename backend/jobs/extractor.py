import os
import json
import traceback
import mammoth
import docx
from utils.storage import resolve_storage_path
from utils.observability import log_event
from utils.diffing import coerce_blocks, compute_block_diff, compute_visual_html_diff

def extract_docx_blocks(filepath):
    """
    Extracts paragraphs and tables from a docx file into a list of block dicts.
    Now tracks headers to provide hierarchical context to the AI.
    """
    doc = docx.Document(filepath)
    blocks = []
    
    # Track current section context
    current_section = "Intro"
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
            
        # Detect if this is a header to update context
        # Most DOCX headers use 'Heading 1', 'Heading 2', etc.
        style_name = para.style.name if para.style else ""
        if style_name.startswith('Heading'):
            current_section = para.text.strip()
            
        blocks.append({
            'type': 'header' if style_name.startswith('Heading') else 'paragraph',
            'text': para.text.strip(),
            'section': current_section
        })

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if any(any(cell for cell in row) for row in table_data):
            blocks.append({
                'type': 'table',
                'data': table_data,
                'text': '\n'.join([' | '.join(row) for row in table_data]),
                'section': current_section
            })
    return blocks

def process_version(app, version_id):
    """Background job to process a document version."""
    with app.app_context():
        from models import db, Version
        version = Version.query.get(version_id)
        if not version:
            return
            
        try:
            log_event(app.logger, "extract_started", version_id=version.id, document_id=version.document_id)
            version.status = 'processing'
            db.session.commit()
            
            storage_root = app.config['STORAGE_ROOT']
            filepath = resolve_storage_path(storage_root, version.storage_path)
            
            # Extract HTML with mammoth
            with open(filepath, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                extracted_html = result.value
                
            # Extract blocks with python-docx
            blocks = extract_docx_blocks(filepath)
            extracted_text = "\n\n".join([b.get('text', '') for b in blocks])
            
            version.extracted_html = extracted_html
            version.extracted_blocks_json = json.dumps(blocks)
            version.extracted_text = extracted_text
            
            # Diffing
            if version.version_number > 1:
                prev_version = Version.query.filter_by(
                    document_id=version.document_id,
                    version_number=version.version_number - 1
                ).first()
                
                if prev_version and prev_version.extracted_text:
                    previous_blocks = coerce_blocks(
                        blocks_json=prev_version.extracted_blocks_json,
                        extracted_text=prev_version.extracted_text
                    )
                    # We use compute_block_diff for stats and AI-ready JSON
                    _, stats, changes = compute_block_diff(previous_blocks, blocks)
                    
                    # We use compute_visual_html_diff for the high-fidelity UI view
                    version.diff_html = compute_visual_html_diff(
                        prev_version.extracted_html, 
                        version.extracted_html
                    )
                    
                    version.stats_json = json.dumps(stats)
                    version.diff_json = json.dumps(changes)
                    log_event(
                        app.logger,
                        "diff_computed",
                        version_id=version.id,
                        previous_version_id=prev_version.id,
                        added_chars=stats.get('added_chars', 0),
                        removed_chars=stats.get('removed_chars', 0),
                        added_blocks=stats.get('added_blocks', 0),
                        removed_blocks=stats.get('removed_blocks', 0),
                        modified_blocks=stats.get('modified_blocks', 0)
                    )
            
            version.status = 'success'
            db.session.commit()
            log_event(app.logger, "extract_succeeded", version_id=version.id, document_id=version.document_id)
            
        except Exception as e:
            version.status = 'failed'
            version.error_message = str(e) + "\\n" + traceback.format_exc()
            db.session.commit()
            log_event(
                app.logger,
                "extract_failed",
                level="error",
                version_id=version.id,
                document_id=version.document_id,
                error=str(e)
            )
